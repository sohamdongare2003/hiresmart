import re
from pathlib import Path


def _extract_pdf_pymupdf(path: str) -> str:
    import fitz
    text_parts = []
    with fitz.open(path) as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)


def _extract_pdf_pdfplumber(path: str) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def extract_pdf_text(path: str) -> str:
    try:
        text = _extract_pdf_pymupdf(path)
        if len(text.strip()) > 50:
            return text
        return _extract_pdf_pdfplumber(path)
    except Exception:
        try:
            return _extract_pdf_pdfplumber(path)
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {e}")


def extract_docx_text(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def extract_doc_text(path: str) -> str:
    import subprocess
    try:
        result = subprocess.run(["antiword", path], capture_output=True, text=True, timeout=15)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    try:
        return extract_docx_text(path)
    except Exception as e:
        raise RuntimeError(f"DOC extraction failed: {e}")


def clean_text(text: str) -> str:
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def extract_text_from_resume(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume not found: {file_path}")
    ext = path.suffix.lower()
    if ext == ".pdf":
        raw = extract_pdf_text(str(path))
    elif ext == ".docx":
        raw = extract_docx_text(str(path))
    elif ext == ".doc":
        raw = extract_doc_text(str(path))
    else:
        raise ValueError(f"Unsupported file type: '{ext}'")
    cleaned = clean_text(raw)
    if not cleaned:
        raise RuntimeError("No text could be extracted from this resume.")
    return cleaned

def extract_candidate_data(text: str) -> dict:

    lines = text.split("\n")

    name = "Unknown Candidate"

    for line in lines[:5]:
        line = line.strip()

        if (
            len(line.split()) >= 2
            and len(line) < 50
            and not "resume" in line.lower()
        ):
            name = line
            break

    email_match = re.search(
        r'[\w\.-]+@[\w\.-]+\.\w+',
        text
    )

    phone_match = re.search(
        r'(\+91[- ]?)?[6-9]\d{9}',
        text
    )

    skill_keywords = [
        "Python",
        "Java",
        "FastAPI",
        "React",
        "MySQL",
        "Django",
        "JavaScript",
        "HTML",
        "CSS",
        "Node.js",
        "MongoDB",
        "SQL",
        "Git",
        "AWS"
    ]

    found_skills = []

    lower_text = text.lower()

    for skill in skill_keywords:
        if skill.lower() in lower_text:
            found_skills.append(skill)

    return {
        "name": name,
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
        "skills": ", ".join(found_skills),
        "experience": "Fresher"
    }